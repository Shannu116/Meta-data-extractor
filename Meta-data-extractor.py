import os
import sys
import stat
import time
import hashlib
import mimetypes
import exif
from PIL import Image
from PIL.ExifTags import TAGS
import exifread
from PyPDF2 import PdfReader
from docx import Document
import datetime
import platform
import gradio as gr
import numpy as np
from scipy.fftpack import dct
import cv2
import io

class WriteBlockedFile(io.IOBase):
    def __init__(self, file_path, mode='rb'):
        if 'w' in mode or 'a' in mode or '+' in mode:
            raise IOError("Write operations are not permitted.")
        self.file = open(file_path, mode)

    def read(self, size=-1):
        return self.file.read(size)

    def readline(self, size=-1):
        return self.file.readline(size)

    def seek(self, offset, whence=io.SEEK_SET):
        return self.file.seek(offset, whence)

    def tell(self):
        return self.file.tell()

    def __enter__(self):
        return self

    def __exit__(self, exc_type, exc_val, exc_tb):
        self.close()

    def close(self):
        self.file.close()

def safe_open(file_path, mode='rb'):
    return WriteBlockedFile(file_path, mode)


def calculate_file_hash(file_path, hash_algorithm="sha256"):
    try:
        with open(file_path, 'rb') as f:
            file_data = f.read()
            hash_obj = hashlib.new(hash_algorithm, file_data)
            file_hash = hash_obj.hexdigest()
            return file_hash
    except FileNotFoundError:
        return f"Error: File '{file_path}' not found."

def get_file_type(file_path):
    mime_type, _ = mimetypes.guess_type(file_path)
    if mime_type:
        return mime_type
    else:
        _, ext = os.path.splitext(file_path)
        ext = ext.lower()
        if ext in ['.jpg', '.jpeg', '.png', '.gif', '.bmp']:
            return 'image/' + ext[1:]
        elif ext == '.pdf':
            return 'application/pdf'
        elif ext == '.docx':
            return 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        else:
            return 'application/octet-stream'

def get_file_metadata(file_path):
    metadata = {}

    # Ensure that the file exists
    if not os.path.exists(file_path):
        metadata['Error'] = f"Error: {file_path} not found"
        return metadata

    try:
        # Get file stats without opening the file
        sb = os.stat(file_path)

        # File type detection
        metadata['File type'] = "unknown"
        mode = sb.st_mode
        if stat.S_ISBLK(mode):
            metadata['File type'] = "block device"
        elif stat.S_ISCHR(mode):
            metadata['File type'] = "character device"
        elif stat.S_ISDIR(mode):
            metadata['File type'] = "directory"
        elif stat.S_ISFIFO(mode):
            metadata['File type'] = "FIFO/pipe"
        elif stat.S_ISLNK(mode):
            metadata['File type'] = "symlink"
        elif stat.S_ISREG(mode):
            metadata['File type'] = "regular file"
        elif stat.S_ISSOCK(mode):
            metadata['File type'] = "socket"

        # Other file information
        metadata['I-node number'] = sb.st_ino
        metadata['Mode'] = oct(sb.st_mode)
        metadata['Link count'] = sb.st_nlink

        # Ownership (UID and GID are not meaningful in Windows)
        if platform.system() != 'Windows':
            metadata['Ownership'] = f"UID={sb.st_uid} GID={sb.st_gid}"
        else:
            metadata['Ownership'] = "(N/A on Windows)"

        # Preferred I/O block size (only available on Unix-like systems)
        if hasattr(sb, 'st_blksize'):
            metadata['Preferred I/O block size'] = f"{sb.st_blksize} bytes"
        else:
            metadata['Preferred I/O block size'] = "(N/A on Windows)"

        # File size and blocks
        metadata['File size'] = f"{sb.st_size} bytes"

        # Blocks allocated (only available on Unix-like systems)
        if hasattr(sb, 'st_blocks'):
            metadata['Blocks allocated'] = sb.st_blocks
        else:
            metadata['Blocks allocated'] = "(N/A on Windows)"

        # File timestamps
        metadata['Last status change'] = time.ctime(sb.st_ctime)
        metadata['Last file access'] = time.ctime(sb.st_atime)
        metadata['Last file modification'] = time.ctime(sb.st_mtime)

        # Additional metadata
        metadata['File Name'] = os.path.basename(file_path)
        metadata['MIME Type'] = mimetypes.guess_type(file_path)[0] or "unknown"
        
        # Calculate hash using write-blocked file access
        metadata['SHA256 Hash'] = calculate_file_hash(file_path)

        # Specific file type metadata
        if metadata['File type'] == "regular file":
            if metadata['MIME Type'].startswith('image'):
                metadata.update(get_image_metadata(file_path))
            elif metadata['MIME Type'] == 'application/pdf':
                metadata.update(get_pdf_metadata(file_path))
            elif metadata['MIME Type'] == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
                metadata.update(get_docx_metadata(file_path))

    except Exception as e:
        metadata['Error'] = str(e)

    return metadata

def calculate_file_hash(file_path, hash_algorithm="sha256"):
    try:
        hash_obj = hashlib.new(hash_algorithm)
        with safe_open(file_path, 'rb') as f:
            for chunk in iter(lambda: f.read(4096), b''):
                hash_obj.update(chunk)
        return hash_obj.hexdigest()
    except Exception as e:
        return f"Error calculating hash: {str(e)}"

def get_image_metadata(file_path):
    image_metadata = {}
    try:
        # Determine image type
        with Image.open(file_path) as img:
            image_format = img.format
            image_metadata['Image Format'] = image_format
            image_metadata['Resolution'] = f"{img.width}x{img.height}"
            image_metadata['Color Mode'] = img.mode

        # EXIF data using exif library (mainly for JPEG)
        if image_format in ['JPEG', 'TIFF']:
            try:
                with open(file_path, 'rb') as img_file:
                    img = exif.Image(img_file)
                    if img.has_exif:
                        image_metadata['Camera Make'] = img.get('make', 'N/A')
                        image_metadata['Camera Model'] = img.get('model', 'N/A')
                        image_metadata['Date Taken'] = img.get('datetime', 'N/A')
                        image_metadata['GPS Info'] = f"{img.get('gps_latitude', 'N/A')}, {img.get('gps_longitude', 'N/A')}"
            except Exception as e:
                image_metadata['EXIF Error'] = str(e)

        # Additional metadata using PIL
        try:
            with Image.open(file_path) as img:
                pil_info = {TAGS.get(tag, tag): str(value) for tag, value in img.getexif().items()}
                image_metadata.update(pil_info)
        except Exception as e:
            image_metadata['PIL Error'] = str(e)
        
        

        # More detailed EXIF data using exifread
        try:
            with open(file_path, 'rb') as f:
                tags = exifread.process_file(f)
                exif_data = {tag: str(tags[tag]) for tag in tags if tag not in ('JPEGThumbnail', 'TIFFThumbnail', 'Filename', 'EXIF MakerNote')}
                image_metadata.update(exif_data)
        except Exception as e:
            image_metadata['ExifRead Error'] = str(e)

        # Analyze for potential Photoshop editing
        software_used = image_metadata.get('Software', 'Unknown')
        if 'Adobe Photoshop' in software_used:
            image_metadata['Forensic Note'] = "WARNING: Image may have been edited with Adobe Photoshop"
        stego_results = detect_steganography(file_path)
        image_metadata.update(stego_results)

    except Exception as e:
        image_metadata['Error'] = f"General error in metadata extraction: {str(e)}"
    try:
        with safe_open(file_path, 'rb') as img_file:
            img = Image.open(img_file)
            image_metadata['Image Format'] = img.format
            image_metadata['Resolution'] = f"{img.width}x{img.height}"
            image_metadata['Color Mode'] = img.mode

            # Extract EXIF data
            exif_data = img._getexif()
            if exif_data:
                for tag_id, value in exif_data.items():
                    tag = TAGS.get(tag_id, tag_id)
                    image_metadata[tag] = str(value)

        # Additional metadata using exifread
        with safe_open(file_path, 'rb') as f:
            tags = exifread.process_file(f)
            for tag, value in tags.items():
                if tag not in ('JPEGThumbnail', 'TIFFThumbnail', 'Filename', 'EXIF MakerNote'):
                    image_metadata[tag] = str(value)

    except Exception as e:
        image_metadata['Error'] = str(e)

    return image_metadata
def detect_steganography(file_path):
    stego_results = {}
    
    # Algorithm 1: LSB (Least Significant Bit) Analysis
    stego_results['LSB Analysis'] = lsb_analysis(file_path)
    
    # Algorithm 2: Chi-Square Analysis
    stego_results['Chi-Square Analysis'] = chi_square_analysis(file_path)
    
    # Algorithm 3: DCT (Discrete Cosine Transform) Analysis
    stego_results['DCT Analysis'] = dct_analysis(file_path)
    
    stego_results['Sample Pair Analysis'] = sample_pair_analysis(file_path)
    stego_results['RS Analysis'] = rs_analysis(file_path)
    stego_results['PVD Analysis'] = pvd_analysis(file_path)
    return stego_results
def lsb_analysis(file_path):
    try:
        img = cv2.imread(file_path)
        if img is None:
            return "Unable to read image file"
        
        lsb = img[:,:,0] % 2
        unusual_patterns = np.sum(lsb) / (img.shape[0] * img.shape[1])
        
        if unusual_patterns > 0.45 and unusual_patterns < 0.55:
            return "Suspicious: Possible LSB steganography detected"
        else:
            return "No obvious LSB steganography detected"
    except Exception as e:
        return f"Error in LSB analysis: {str(e)}"
def chi_square_analysis(file_path):
    try:
        img = cv2.imread(file_path, 0)  # Read as grayscale
        if img is None:
            return "Unable to read image file"
        
        hist = cv2.calcHist([img], [0], None, [256], [0, 256])
        even_hist = hist[::2]
        odd_hist = hist[1::2]
        
        chi_square = np.sum((even_hist - odd_hist)**2 / (even_hist + odd_hist + 1e-6))
        
        if chi_square < 0.1:
            return "Suspicious: Possible steganography detected by Chi-Square analysis"
        else:
            return "No steganography detected by Chi-Square analysis"
    except Exception as e:
        return f"Error in Chi-Square analysis: {str(e)}"

def dct_analysis(file_path):
    try:
        img = cv2.imread(file_path, 0)  # Read as grayscale
        if img is None:
            return "Unable to read image file"
        
        img_dct = dct(dct(img.T, norm='ortho').T, norm='ortho')
        dct_values = np.abs(img_dct.flatten())
        
        threshold = np.percentile(dct_values, 99.95)
        suspicious_coeffs = np.sum(dct_values > threshold)
        
        if suspicious_coeffs > 100:
            return "Suspicious: Possible steganography detected by DCT analysis"
        else:
            return "No obvious steganography detected by DCT analysis"
    except Exception as e:
        return f"Error in DCT analysis: {str(e)}"

def sample_pair_analysis(file_path):
    try:
        img = cv2.imread(file_path, 0)  # Read as grayscale
        if img is None:
            return "Unable to read image file"
        
        rows, cols = img.shape
        pairs = np.column_stack((img[:-1, :].flatten(), img[1:, :].flatten()))
        
        even_pairs = pairs[np.sum(pairs % 2, axis=1) == 0]
        odd_pairs = pairs[np.sum(pairs % 2, axis=1) == 1]
        
        beta = len(even_pairs) / (len(even_pairs) + len(odd_pairs))
        
        if 0.45 < beta < 0.55:
            return "Suspicious: Possible steganography detected by Sample Pair Analysis"
        else:
            return "No steganography detected by Sample Pair Analysis"
    except Exception as e:
        return f"Error in Sample Pair Analysis: {str(e)}"

def rs_analysis(file_path):
    try:
        img = cv2.imread(file_path, 0)  # Read as grayscale
        if img is None:
            return "Unable to read image file"
        
        def flip_lsb(x):
            return x ^ 1
        
        rows, cols = img.shape
        
        r_m, s_m, r_m_inv, s_m_inv = 0, 0, 0, 0
        
        for i in range(0, rows-1, 2):
            for j in range(0, cols-1, 2):
                block = img[i:i+2, j:j+2].astype(int)
                flipped_block = np.vectorize(flip_lsb)(block)
                
                diff = np.sum(np.abs(block - np.roll(block, 1, axis=1)))
                diff_flipped = np.sum(np.abs(flipped_block - np.roll(flipped_block, 1, axis=1)))
                
                if diff < diff_flipped:
                    r_m += 1
                elif diff > diff_flipped:
                    s_m += 1
                
                diff_inv = np.sum(np.abs(block - np.roll(block, -1, axis=1)))
                diff_flipped_inv = np.sum(np.abs(flipped_block - np.roll(flipped_block, -1, axis=1)))
                
                if diff_inv < diff_flipped_inv:
                    r_m_inv += 1
                elif diff_inv > diff_flipped_inv:
                    s_m_inv += 1
        
        total = r_m + s_m
        total_inv = r_m_inv + s_m_inv
        
        if total == 0 or total_inv == 0:
            return "Unable to perform RS Analysis: division by zero"
        
        d = abs((r_m - s_m) / total) - abs((r_m_inv - s_m_inv) / total_inv)
        
        if abs(d) < 0.05:
            return "Suspicious: Possible steganography detected by RS Analysis"
        else:
            return "No steganography detected by RS Analysis"
    except Exception as e:
        return f"Error in RS Analysis: {str(e)}"
def pvd_analysis(file_path):
    try:
        img = cv2.imread(file_path, 0)  # Read as grayscale
        if img is None:
            return "Unable to read image file"
        
        rows, cols = img.shape
        diff_hist = np.zeros(256)
        
        for i in range(rows):
            for j in range(cols - 1):
                diff = abs(int(img[i, j]) - int(img[i, j+1]))
                diff_hist[diff] += 1
        
        # Normalize the histogram
        diff_hist /= np.sum(diff_hist)
        
        # Check for unusual peaks in the histogram
        peaks = np.where(diff_hist > np.mean(diff_hist) + 2 * np.std(diff_hist))[0]
        
        if len(peaks) > 5:  # Threshold can be adjusted
            return "Suspicious: Possible steganography detected by PVD Analysis"
        else:
            return "No obvious steganography detected by PVD Analysis"
    except Exception as e:
        return f"Error in PVD Analysis: {str(e)}"
def get_pdf_metadata(file_path):
    pdf_metadata = {}
    try:
        with open(file_path, 'rb') as pdf_file:
            pdf = PdfReader(pdf_file)
            info = pdf.metadata
            pdf_metadata['Author'] = info.author if info.author else 'N/A'
            pdf_metadata['Creator'] = info.creator if info.creator else 'N/A'
            pdf_metadata['Producer'] = info.producer if info.producer else 'N/A'
            pdf_metadata['Subject'] = info.subject if info.subject else 'N/A'
            pdf_metadata['Title'] = info.title if info.title else 'N/A'
            pdf_metadata['Creation Date'] = info.creation_date if info.creation_date else 'N/A'
            pdf_metadata['Modification Date'] = info.modification_date if info.modification_date else 'N/A'
            pdf_metadata['Number of Pages'] = len(pdf.pages)
    except Exception as e:
        pdf_metadata['Error'] = str(e)
    
    try:
        with safe_open(file_path, 'rb') as pdf_file:
            pdf = PdfReader(pdf_file)
            info = pdf.metadata
            if info:
                for key, value in info.items():
                    pdf_metadata[key] = str(value)
            pdf_metadata['Number of Pages'] = len(pdf.pages)
    except Exception as e:
        pdf_metadata['Error'] = str(e)
    return pdf_metadata

def get_docx_metadata(file_path):
    docx_metadata = {}
    try:
        doc = Document(file_path)
        core_properties = doc.core_properties
        docx_metadata['Author'] = core_properties.author if core_properties.author else 'N/A'
        docx_metadata['Created'] = core_properties.created.strftime('%Y-%m-%d %H:%M:%S') if core_properties.created else 'N/A'
        docx_metadata['Last Modified By'] = core_properties.last_modified_by if core_properties.last_modified_by else 'N/A'
        docx_metadata['Last Printed'] = core_properties.last_printed.strftime('%Y-%m-%d %H:%M:%S') if core_properties.last_printed else 'N/A'
        docx_metadata['Title'] = core_properties.title if core_properties.title else 'N/A'
        docx_metadata['Subject'] = core_properties.subject if core_properties.subject else 'N/A'
        docx_metadata['Keywords'] = core_properties.keywords if core_properties.keywords else 'N/A'
        docx_metadata['Comments'] = core_properties.comments if core_properties.comments else 'N/A'
        docx_metadata['Category'] = core_properties.category if core_properties.category else 'N/A'
        docx_metadata['Number of Pages'] = doc.element.body.get_page_count()
    except Exception as e:
        docx_metadata['Error'] = str(e)
    try:
        with safe_open(file_path, 'rb') as docx_file:
            doc = Document(docx_file)
            core_properties = doc.core_properties
            docx_metadata['Author'] = core_properties.author
            docx_metadata['Created'] = str(core_properties.created)
            docx_metadata['Last Modified By'] = core_properties.last_modified_by
            docx_metadata['Last Printed'] = str(core_properties.last_printed)
            docx_metadata['Title'] = core_properties.title
            docx_metadata['Subject'] = core_properties.subject
            docx_metadata['Keywords'] = core_properties.keywords
            docx_metadata['Comments'] = core_properties.comments
            docx_metadata['Category'] = core_properties.category
            docx_metadata['Number of Pages'] = len(doc.paragraphs)
    except Exception as e:
        docx_metadata['Error'] = str(e)
    return docx_metadata

def extract_metadata(file):
    if file is None:
        return "No file uploaded. Please upload a file to extract metadata."
    
    file_path = file.name
    metadata = get_file_metadata(file_path)
    
    result = "\nComprehensive File Metadata Analysis Results:\n"
    result += "=============================================\n"
    for key, value in metadata.items():
        if isinstance(value, dict):
            result += f"{key}:\n"
            for sub_key, sub_value in value.items():
                result += f"  {sub_key}: {sub_value}\n"
        else:
            result += f"{key}: {value}\n"
    
    return result

iface = gr.Interface(
    fn=extract_metadata,
    inputs=gr.File(label="Upload File for Comprehensive Analysis"),
    outputs=gr.Textbox(label="Comprehensive File Metadata Analysis Results", lines=30),
    title="Comprehensive File Metadata Analyzer",
    description="Upload a file for detailed metadata analysis. This tool extracts extensive metadata, including file system attributes, timestamps, and file-type specific information."
)

if __name__ == "__main__":
    iface.launch()
